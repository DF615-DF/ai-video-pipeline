import os
import re
import shutil

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


BASE = r"E:\AI\工作\my-knowledge\01-项目\缘结守"
OUT = r"E:\AI\工作\缘结守交付"
ZIP_BASE = r"E:\AI\工作\缘结守-交付"
OUT_STORYLINE = r"E:\AI\工作\缘结守-十幕故事线交付"
ZIP_STORYLINE = r"E:\AI\工作\缘结守-十幕故事线"


def read(name):
    with open(os.path.join(BASE, name), encoding="utf-8") as f:
        return f.read()


def configure_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "微软雅黑")


def clean_text(text):
    text = re.sub(r"!\[\[([^]]+)\]\]", r"\1", text)
    text = re.sub(r"\[\[([^]]+)\]\]", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text


def is_table_separator(cells):
    pattern = re.compile(r"^:?-{3,}:?$")
    return len(cells) > 0 and all(pattern.match(cell.strip()) for cell in cells)


def add_markdown(doc, markdown, base_level=1):
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)) + base_level - 1, 9)
            doc.add_heading(clean_text(heading.group(2)), level=level)
            i += 1
            continue

        if line.startswith(">"):
            doc.add_paragraph(clean_text(line.lstrip("> ")))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            rows = []
            for table_line in table_lines:
                cells = [cell.strip() for cell in table_line.strip("|").split("|")]
                if is_table_separator(cells):
                    continue
                rows.append(cells)

            if rows:
                cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for col_index in range(cols):
                        text = row[col_index] if col_index < len(row) else ""
                        table.cell(row_index, col_index).text = clean_text(text)
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(clean_text(stripped[2:]), style="List Bullet")
            i += 1
            continue

        image = re.match(r"^!\[\[([^]]+)\]\]$", stripped)
        if image:
            image_path = os.path.join(BASE, image.group(1))
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(3.5))
            i += 1
            continue

        if stripped.startswith("---"):
            i += 1
            continue

        doc.add_paragraph(clean_text(line))
        i += 1


def build_setting_doc():
    doc = Document()
    configure_document(doc)
    doc.add_heading("缘结守 - 剧情设定简介", 0)

    add_markdown(doc, read("缘结守-故事简介.md"), base_level=1)

    doc.add_page_break()
    add_markdown(doc, read("缘结守-角色阶段设定.md"), base_level=1)

    doc.add_page_break()
    add_markdown(doc, read("缘结守-项目初稿.md"), base_level=1)

    path = os.path.join(OUT, "缘结守-剧情设定简介.docx")
    doc.save(path)
    return path


def build_storyboard_doc():
    doc = Document()
    configure_document(doc)
    doc.add_heading("缘结守 - 十幕分镜细稿", 0)

    add_markdown(doc, read("缘结守-全片节奏总览.md"), base_level=1)

    act_names = [
        ("第一幕", "雪夜神龛", "缘结守-第一幕细稿.md"),
        ("第二幕", "春日田埂", "缘结守-第二幕细稿.md"),
        ("第三幕", "夏日溪边", "缘结守-第三幕细稿.md"),
        ("第四幕", "校园守护", "缘结守-第四幕细稿.md"),
        ("第五幕", "夏夜烟花", "缘结守-第五幕细稿.md"),
        ("第六幕", "车站分离", "缘结守-第六幕细稿.md"),
        ("第七幕", "东京思念", "缘结守-第七幕细稿.md"),
        ("第八幕", "枯木异变", "缘结守-第八幕细稿.md"),
        ("第九幕", "挡下施工", "缘结守-第九幕细稿.md"),
        ("第十幕", "春日神龛", "缘结守-第十幕细稿.md"),
    ]

    for index, (act, title, filename) in enumerate(act_names, start=1):
        doc.add_page_break()
        doc.add_heading(f"{act}：{title}", level=1)
        add_markdown(doc, read(filename), base_level=2)

        if index == 1:
            doc.add_page_break()
            add_markdown(doc, read("缘结守-第一幕分镜.md"), base_level=2)

    path = os.path.join(OUT, "缘结守-十幕分镜细稿.docx")
    doc.save(path)
    return path


def build_storyline_doc():
    doc = Document()
    configure_document(doc)
    doc.add_heading("缘结守 - 十幕故事线", 0)

    acts = [
        ("第一幕", "雪夜神龛", "0-7 分钟", "缘结守-第一幕细稿.md"),
        ("第二幕", "春日田埂", "7-14 分钟", "缘结守-第二幕细稿.md"),
        ("第三幕", "夏日溪边", "14-21 分钟", "缘结守-第三幕细稿.md"),
        ("第四幕", "校园守护", "21-28 分钟", "缘结守-第四幕细稿.md"),
        ("第五幕", "夏夜烟花", "28-36 分钟", "缘结守-第五幕细稿.md"),
        ("第六幕", "车站分离", "36-43 分钟", "缘结守-第六幕细稿.md"),
        ("第七幕", "东京思念", "43-54 分钟", "缘结守-第七幕细稿.md"),
        ("第八幕", "枯木异变", "54-62 分钟", "缘结守-第八幕细稿.md"),
        ("第九幕", "挡下施工", "62-70 分钟", "缘结守-第九幕细稿.md"),
        ("第十幕", "春日神龛", "70-80 分钟", "缘结守-第十幕细稿.md"),
    ]

    table = doc.add_table(rows=len(acts) + 1, cols=3)
    table.style = "Table Grid"
    headers = ["幕", "名称", "时间"]
    for col, header in enumerate(headers):
        table.cell(0, col).text = header
    for row, (act, title, duration, _) in enumerate(acts, start=1):
        table.cell(row, 0).text = act
        table.cell(row, 1).text = title
        table.cell(row, 2).text = duration

    doc.add_page_break()
    add_markdown(doc, read("缘结守-故事简介.md"), base_level=1)

    for act, title, duration, filename in acts:
        doc.add_page_break()
        doc.add_heading(f"{act}：{title}（{duration}）", level=1)
        add_markdown(doc, read(filename), base_level=2)

    os.makedirs(OUT_STORYLINE, exist_ok=True)
    path = os.path.join(OUT_STORYLINE, "缘结守-十幕故事线.docx")
    doc.save(path)
    return path


def copy_images():
    for filename in os.listdir(BASE):
        if filename.lower().endswith((".jpg", ".jpeg", ".png")):
            shutil.copy2(os.path.join(BASE, filename), os.path.join(OUT, filename))


def main():
    storyline_path = build_storyline_doc()
    zip_path = shutil.make_archive(ZIP_STORYLINE, "zip", OUT_STORYLINE)
    print(storyline_path)
    print(zip_path)


if __name__ == "__main__":
    main()
